"""HTTP-level tests for the Writing Review backend routes.

Every test spins up an aiohttp application registered by
:func:`register_routes`, injects a fake provider, and hits the endpoints
via ``AioHTTPTestCase``'s async ``self.client``.

* :class:`TestScanRoute`           -> POST /scan validation and job dispatch
* :class:`TestJobStatusRoute`      -> GET /jobs/{job_id}
* :class:`TestReviewsRoutes`       -> GET/DELETE /reviews and /reviews/{id}
* :class:`TestSettingsRoutes`      -> GET/PATCH /settings
"""

from __future__ import annotations

import asyncio
import json
import os
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

import aiohttp
from aiohttp import web
from aiohttp.test_utils import AioHTTPTestCase

from kiro_crew.apps.builtins.writing_review import (
    ReviewContext,
    ScanResult,
    Section,
)
from kiro_crew.apps.builtins.writing_review.backend import routes as routes_module
from kiro_crew.apps.builtins.writing_review.backend.routes import (
    _reset_jobs_for_tests,
    register_routes,
)
from kiro_crew.apps.builtins.writing_review.backend.store import save_review


class _WritingReviewRoutesTestBase(AioHTTPTestCase):
    """Common setup: pin KIROCREW_HOME, register routes, inject a fake provider."""

    async def get_application(self) -> web.Application:
        application = web.Application()
        register_routes(application)
        return application

    async def asyncSetUp(self) -> None:
        self._tempdir = TemporaryDirectory()
        self._previous_home = os.environ.get("KIROCREW_HOME")
        os.environ["KIROCREW_HOME"] = self._tempdir.name
        _reset_jobs_for_tests()
        await super().asyncSetUp()

    async def asyncTearDown(self) -> None:
        await super().asyncTearDown()
        if self._previous_home is None:
            os.environ.pop("KIROCREW_HOME", None)
        else:
            os.environ["KIROCREW_HOME"] = self._previous_home
        self._tempdir.cleanup()


class TestScanRoute(_WritingReviewRoutesTestBase):
    """POST /scan input validation and background dispatch."""

    async def test_scan_endpoint_rejects_missing_doc_input(self) -> None:
        response = await self.client.post("/api/apps/writing-review/scan", json={})

        self.assertEqual(response.status, 400)

    async def test_scan_endpoint_returns_job_id_when_valid(self) -> None:
        document_path = Path(self._tempdir.name) / "example.md"
        document_path.write_text("# Intro\nHello\n", encoding="utf-8")

        fake_scan_result = ScanResult(
            doc_path=str(document_path),
            doc_name="example.md",
            doc_context=ReviewContext(),
            sections=[Section(heading="Intro", body="Hello")],
            findings=[],
            verdict="green",
            scanners_run=["clarity"],
        )

        async def fake_run_scan(**_kwargs):
            return fake_scan_result

        with patch.object(routes_module, "run_scan", fake_run_scan):
            response = await self.client.post(
                "/api/apps/writing-review/scan",
                json={
                    "doc_path": str(document_path),
                    "context": {
                        "audience": "team",
                        "doc_type": "update",
                        "tone": "neutral",
                    },
                },
            )

            self.assertEqual(response.status, 200)
            body = await response.json()
            self.assertIn("job_id", body)

            # Let the background task settle before we tear down the loop.
            await asyncio.sleep(0.05)

    async def test_scan_endpoint_rejects_bad_scanner_toggles(self) -> None:
        response = await self.client.post(
            "/api/apps/writing-review/scan",
            json={
                "doc_path": "/tmp/x.md",
                "scanner_toggles": "not-an-object",
            },
        )
        self.assertEqual(response.status, 400)

    async def test_scan_endpoint_forwards_ask_into_run_scan_context(self) -> None:
        # The frontend NewReviewDialog will POST ``context.ask`` on submit;
        # the scan handler MUST parse it out of the payload and build a
        # ``ReviewContext`` with the ask field populated so downstream
        # scanner prompts and the discussion-agent context bundle can
        # surface it. Regression guard against a merge that adds an
        # ``ask`` textarea to the dialog but forgets to plumb it here.
        document_path = Path(self._tempdir.name) / "example.md"
        document_path.write_text("# Intro\nHello\n", encoding="utf-8")

        captured_contexts: list[ReviewContext] = []

        async def fake_run_scan(**kwargs: object) -> ScanResult:
            passed_context = kwargs.get("context")
            assert isinstance(passed_context, ReviewContext)
            captured_contexts.append(passed_context)
            return ScanResult(
                doc_path=str(document_path),
                doc_name="example.md",
                doc_context=passed_context,
                sections=[Section(heading="Intro", body="Hello")],
                findings=[],
                verdict="green",
                scanners_run=["clarity"],
            )

        with patch.object(routes_module, "run_scan", fake_run_scan):
            response = await self.client.post(
                "/api/apps/writing-review/scan",
                json={
                    "doc_path": str(document_path),
                    "context": {
                        "audience": "team",
                        "doc_type": "update",
                        "tone": "neutral",
                        "ask": "Focus on whether the rollout timeline is realistic.",
                    },
                },
            )
            self.assertEqual(response.status, 200)
            # Give the background scan task a chance to execute and land
            # its call into ``fake_run_scan``. The response returned as
            # soon as the task was scheduled; the capture happens later.
            for _ in range(20):
                if captured_contexts:
                    break
                await asyncio.sleep(0.02)

        self.assertEqual(len(captured_contexts), 1)
        self.assertEqual(
            captured_contexts[0].ask,
            "Focus on whether the rollout timeline is realistic.",
        )


class TestJobStatusRoute(_WritingReviewRoutesTestBase):
    """GET /jobs/{job_id} returns 404 for unknown jobs and 200 for known ones."""

    async def test_missing_job_returns_404(self) -> None:
        response = await self.client.get("/api/apps/writing-review/jobs/nonexistent")
        self.assertEqual(response.status, 404)

    async def test_known_job_returns_current_state(self) -> None:
        await routes_module._record_job_state("abc123", status="running", phase="starting")

        response = await self.client.get("/api/apps/writing-review/jobs/abc123")

        self.assertEqual(response.status, 200)
        body = await response.json()
        self.assertEqual(body["status"], "running")
        self.assertEqual(body["phase"], "starting")


class TestReviewsRoutes(_WritingReviewRoutesTestBase):
    """GET/DELETE /reviews and /reviews/{id} against a temp data home."""

    async def test_reviews_list_endpoint(self) -> None:
        response = await self.client.get("/api/apps/writing-review/reviews")

        self.assertEqual(response.status, 200)
        body = await response.json()
        self.assertIn("reviews", body)
        self.assertEqual(body["reviews"], [])

    async def test_review_detail_missing_returns_404(self) -> None:
        response = await self.client.get("/api/apps/writing-review/reviews/nonexistent")
        self.assertEqual(response.status, 404)

    async def test_delete_review_missing_returns_404(self) -> None:
        response = await self.client.delete("/api/apps/writing-review/reviews/nonexistent")
        self.assertEqual(response.status, 404)


class TestReviewContextRoute(_WritingReviewRoutesTestBase):
    """Behaviour #10 -- GET /reviews/{id}/context returns the discussion bundle."""

    async def test_context_missing_review_returns_404(self) -> None:
        response = await self.client.get("/api/apps/writing-review/reviews/nonexistent/context")
        self.assertEqual(response.status, 404)

    async def test_context_returns_review_and_document_bundle(self) -> None:
        # Write a document to disk that the endpoint will inline into the response.
        doc_path = Path(self._tempdir.name) / "sample-doc.md"
        doc_path.write_text("# Sample\nBody paragraph.\n", encoding="utf-8")

        # Persist a review pointing at that document.
        scan_result = ScanResult(
            doc_path=str(doc_path),
            doc_name=doc_path.name,
            doc_context=ReviewContext(audience="team"),
            sections=[Section(heading="Sample", body="Body paragraph.")],
            findings=[],
            verdict="green",
            scanners_run=["clarity"],
            partial_failure=False,
            failed_scanners=[],
        )
        review_id = save_review(scan_result)

        response = await self.client.get(f"/api/apps/writing-review/reviews/{review_id}/context")
        self.assertEqual(response.status, 200)
        body = await response.json()
        self.assertIn("review", body)
        self.assertIn("document_content", body)
        self.assertIn("scanner_brief_dir", body)
        self.assertIn("Body paragraph.", body["document_content"])
        self.assertEqual(body["review"]["id"], review_id)

    async def test_context_returns_extracted_prose_for_docx_doc_path(self) -> None:
        # When the review's ``doc_path`` is a binary ``.docx``, the
        # endpoint MUST NOT try to ``read_text`` the raw ZIP bytes --
        # that raises ``UnicodeDecodeError`` (not caught by the
        # OSError clause), 500s the endpoint, and leaves the frontend
        # falling back to the compact ``review_id`` marker so the
        # writing-review-reviewer agent gets no findings. Instead the
        # endpoint MUST route docx through ``parse_doc`` and return
        # the extracted prose so the agent can reason about the
        # document.
        from docx import Document as DocxDocument

        docx_path = Path(self._tempdir.name) / "sample-doc.docx"
        docx_document = DocxDocument()
        docx_document.add_heading("Sample heading", level=1)
        docx_document.add_paragraph("Body paragraph in the docx.")
        docx_document.save(str(docx_path))

        scan_result = ScanResult(
            doc_path=str(docx_path),
            doc_name=docx_path.name,
            doc_context=ReviewContext(audience="team"),
            sections=[Section(heading="Sample heading", body="Body paragraph in the docx.")],
            findings=[],
            verdict="green",
            scanners_run=["clarity"],
            partial_failure=False,
            failed_scanners=[],
        )
        review_id = save_review(scan_result)

        response = await self.client.get(
            f"/api/apps/writing-review/reviews/{review_id}/context",
        )
        self.assertEqual(response.status, 200)
        body = await response.json()
        # ``document_content`` MUST contain the extracted prose. The
        # exact rendering shape is up to the endpoint, but the body
        # text a scanner would have seen MUST be present so the
        # discussion agent can reason about it.
        self.assertIn("Body paragraph in the docx.", body["document_content"])
        self.assertIn("Sample heading", body["document_content"])


class TestSettingsRoutes(_WritingReviewRoutesTestBase):
    """GET/PATCH /settings round-trip user configuration."""

    async def test_settings_get_returns_defaults(self) -> None:
        response = await self.client.get("/api/apps/writing-review/settings")

        self.assertEqual(response.status, 200)
        body = await response.json()
        self.assertIn("scanner_toggles", body)
        self.assertTrue(body["scanner_toggles"]["clarity"])

    async def test_settings_patch_merges(self) -> None:
        response = await self.client.patch(
            "/api/apps/writing-review/settings",
            json={"default_audience": "VP"},
        )

        self.assertEqual(response.status, 200)
        body = await response.json()
        self.assertEqual(body["default_audience"], "VP")
        # Defaults preserved.
        self.assertEqual(body["default_tone"], "")


class TestJobStateMonotonicity(_WritingReviewRoutesTestBase):
    """Regression guard for the ``on_phase("done")`` race.

    ``run_scan`` fires an ``on_phase("done", ...)`` callback that this backend
    schedules via ``asyncio.create_task`` with ``status="running"``. That task
    can be scheduled BEFORE, but run AFTER, the driver's own completion write
    (``status="done"``). Without a monotonicity guard on ``_record_job_state``
    the late ``running`` write clobbers the completion — the review is saved
    on disk, but the job record on disk (and in memory) says the job is still
    running, so the frontend loop never terminates and the user sees a stuck
    "Scanning..." card with the completed review sitting next to it.

    The invariant these tests pin: once a job's status is a terminal value
    (``done`` / ``failed`` / ``interrupted``), it cannot be downgraded back
    to ``running`` by any later write.
    """

    async def test_done_state_is_not_downgraded_to_running(self) -> None:
        await routes_module._record_job_state(
            "job-monotonic-done",
            status="done",
            phase="done",
            review_id="review-abc",
        )
        # Simulate the late ``on_phase("done")`` task landing after completion.
        await routes_module._record_job_state(
            "job-monotonic-done",
            status="running",
            phase="done",
        )
        async with routes_module._JOBS_LOCK:
            snapshot = dict(routes_module._JOBS["job-monotonic-done"])
        self.assertEqual(snapshot["status"], "done")
        self.assertEqual(snapshot["review_id"], "review-abc")


class TestJobLoadRecovery(_WritingReviewRoutesTestBase):
    """``_load_jobs_from_disk`` recovery of stuck ``running + phase=done`` records.

    The pre-guard bug produced a specific broken shape on disk: ``status =
    "running"`` with ``phase = "done"`` and ``review_id = null``. The scan
    itself completed (``save_review`` ran before the clobber), so the review
    exists in the reviews store; only the job record is stale. Downgrading
    the whole class to ``interrupted`` would tell the user "re-run this",
    which is misleading — the finished review is right there in the sidebar.
    Recovery downgrades this specific shape to ``done`` so the frontend's
    ``ScanProgress`` useEffect clears ``activeJobId`` cleanly (its ``done``
    branch tolerates a missing ``review_id`` and falls through to the empty
    state, letting the user pick the completed review from the list).
    """

    async def test_running_with_done_phase_is_recovered_as_done(self) -> None:
        jobs_file = routes_module._jobs_file_path()
        jobs_file.parent.mkdir(parents=True, exist_ok=True)
        clobbered_record = {
            "id": "stuck-job",
            "status": "running",
            "phase": "done",
            "detail": {"findings": 42},
            "review_id": None,
            "error": None,
            "doc_name": "hapi.md",
            "updated_at": 1_700_000_000.0,
        }
        jobs_file.write_text(json.dumps([clobbered_record]), encoding="utf-8")

        routes_module._load_jobs_from_disk()

        async with routes_module._JOBS_LOCK:
            recovered = dict(routes_module._JOBS["stuck-job"])
        self.assertEqual(recovered["status"], "done")
        self.assertEqual(recovered["phase"], "done")

    async def test_running_without_done_phase_is_recovered_as_interrupted(self) -> None:
        """Genuinely interrupted scans still downgrade to ``interrupted`` --
        only the ``phase == "done"`` shape is treated as recovered-done.
        """
        jobs_file = routes_module._jobs_file_path()
        jobs_file.parent.mkdir(parents=True, exist_ok=True)
        mid_scan_record = {
            "id": "mid-scan-job",
            "status": "running",
            "phase": "scanner",
            "detail": {},
            "review_id": None,
            "error": None,
            "doc_name": "wip.md",
            "updated_at": 1_700_000_000.0,
        }
        jobs_file.write_text(json.dumps([mid_scan_record]), encoding="utf-8")

        routes_module._load_jobs_from_disk()

        async with routes_module._JOBS_LOCK:
            recovered = dict(routes_module._JOBS["mid-scan-job"])
        self.assertEqual(recovered["status"], "interrupted")


class TestSanitizeDocFilename(unittest.TestCase):
    """``_sanitize_doc_filename`` — filename hygiene for browse-uploaded docs.

    Browse-uploaded docs on the frontend send the original filename so the
    review record can display something human ("hapi_design_doc.md") instead
    of the storage key ("abc12345_pasted.md"). That filename is user-supplied
    string data — never a path — and this sanitiser is the guardrail that
    stops it from being used to attack the uploads directory (path traversal,
    control chars, oversize names).
    """

    def test_plain_filename_survives_unchanged(self) -> None:
        self.assertEqual(
            routes_module._sanitize_doc_filename("hapi_design_doc.md"),
            "hapi_design_doc.md",
        )

    def test_strips_path_traversal_segments(self) -> None:
        # A malicious original_filename must never resolve outside the
        # uploads directory. The sanitiser takes the ``basename`` so
        # embedded slashes / dot-dot segments cannot escape.
        self.assertEqual(
            routes_module._sanitize_doc_filename("../../etc/passwd.md"),
            "passwd.md",
        )
        self.assertEqual(
            routes_module._sanitize_doc_filename("subdir/notes.md"),
            "notes.md",
        )

    def test_strips_windows_path_separators(self) -> None:
        # Browser file inputs on Windows can hand us paths with backslash
        # separators. The sanitiser must strip Windows-style directory
        # prefixes the same way it strips POSIX ones — the on-disk name
        # is the uuid-prefixed storage key, and only the basename is
        # kept for human display in the review record.
        self.assertEqual(
            routes_module._sanitize_doc_filename("C:\\Users\\alice\\design.md"),
            "design.md",
        )
        self.assertEqual(
            routes_module._sanitize_doc_filename("subdir\\notes.md"),
            "notes.md",
        )
        # Mixed separators — a Windows-shell copy-paste can contain both.
        self.assertEqual(
            routes_module._sanitize_doc_filename("subdir/nested\\notes.md"),
            "notes.md",
        )

    def test_replaces_unsafe_characters_with_underscore(self) -> None:
        # Anything outside ``[A-Za-z0-9._-]`` collapses to a single underscore.
        self.assertEqual(
            routes_module._sanitize_doc_filename("my doc (v2).md"),
            "my_doc_v2_.md",
        )

    def test_strips_null_and_control_characters(self) -> None:
        self.assertEqual(
            routes_module._sanitize_doc_filename("note\x00.md"),
            "note_.md",
        )
        self.assertEqual(
            routes_module._sanitize_doc_filename("evil\x1b.md"),
            "evil_.md",
        )

    def test_falls_back_when_result_is_empty(self) -> None:
        # A filename that sanitises to the empty string (e.g. all-emoji) MUST
        # produce a usable fallback rather than an empty string that would
        # collapse into just the uuid prefix on the way to disk.
        self.assertEqual(routes_module._sanitize_doc_filename(""), "pasted.md")
        self.assertEqual(routes_module._sanitize_doc_filename("///"), "pasted.md")

    def test_caps_length_to_prevent_filesystem_limits(self) -> None:
        # Most POSIX filesystems cap a component at 255 bytes. Cap well below
        # that (128) to leave room for the uuid prefix and any future extension.
        long_original = "a" * 500 + ".md"
        sanitized = routes_module._sanitize_doc_filename(long_original)
        self.assertLessEqual(len(sanitized), 128)
        # Extension must be preserved for the file-type router downstream.
        self.assertTrue(sanitized.endswith(".md"))


class TestStashPastedDocument(_WritingReviewRoutesTestBase):
    """``_stash_pasted_document`` includes the sanitised original filename in
    the storage name when one is supplied.

    Storage stays uuid-prefixed for collision safety — two browse-uploads
    of files that happen to share a filename must not clobber each other on
    disk. The original name rides along as the suffix so the on-disk file
    is human-scannable and the ``doc_name`` displayed to the user can be
    derived by stripping the uuid prefix.
    """

    async def test_uses_uuid_pasted_default_when_no_original_provided(self) -> None:
        # Backward compatibility: an existing caller passing no
        # ``original_filename`` MUST keep getting the ``{uuid}_pasted.md``
        # shape it had before this change.
        stored_path = await routes_module._stash_pasted_document("hello world")
        self.assertTrue(stored_path.endswith("_pasted.md"))
        self.assertTrue(Path(stored_path).is_file())

    async def test_includes_sanitized_original_filename_when_provided(self) -> None:
        stored_path = await routes_module._stash_pasted_document(
            "hello world",
            original_filename="hapi_design_doc.md",
        )
        stored_name = Path(stored_path).name
        self.assertTrue(stored_name.endswith("_hapi_design_doc.md"))
        self.assertTrue(Path(stored_path).is_file())

    async def test_applies_filename_sanitisation_before_storage(self) -> None:
        stored_path = await routes_module._stash_pasted_document(
            "content",
            original_filename="../../etc/passwd.md",
        )
        stored_name = Path(stored_path).name
        # Path traversal segments are stripped by the sanitiser — the
        # storage name must NOT contain any ``..`` or slashes.
        self.assertNotIn("..", stored_name)
        self.assertNotIn("/", stored_name)
        self.assertTrue(stored_name.endswith("_passwd.md"))


class TestScanRouteDocName(_WritingReviewRoutesTestBase):
    """POST /scan carries the user-supplied ``doc_name`` through to the job.

    The frontend's browse-file flow POSTs ``doc_text`` alongside a
    ``doc_name`` field carrying the original filename ("hapi_design_doc.md").
    The backend MUST use the sanitised form of that name as the job
    record's ``doc_name`` so the sidebar in-progress card and the
    downstream review record display the human filename rather than the
    ``{uuid}_pasted.md`` storage key.
    """

    async def test_doc_name_from_request_is_stored_on_job_record(self) -> None:
        fake_scan_result = ScanResult(
            doc_path="/tmp/x.md",
            doc_name="ignored",
            doc_context=ReviewContext(),
            sections=[Section(heading="", body="")],
            findings=[],
            verdict="green",
            scanners_run=[],
        )

        async def fake_run_scan(**_kwargs):
            return fake_scan_result

        with patch.object(routes_module, "run_scan", fake_run_scan):
            response = await self.client.post(
                "/api/apps/writing-review/scan",
                json={
                    "doc_text": "hello world",
                    "doc_name": "hapi_design_doc.md",
                },
            )
            self.assertEqual(response.status, 200)
            body = await response.json()
            job_id = body["job_id"]

            async with routes_module._JOBS_LOCK:
                job_record = dict(routes_module._JOBS[job_id])
            self.assertEqual(job_record["doc_name"], "hapi_design_doc.md")

            # Let the background task settle so teardown is clean.
            await asyncio.sleep(0.05)

    async def test_doc_name_is_sanitised_on_entry(self) -> None:
        fake_scan_result = ScanResult(
            doc_path="/tmp/x.md",
            doc_name="ignored",
            doc_context=ReviewContext(),
            sections=[Section(heading="", body="")],
            findings=[],
            verdict="green",
            scanners_run=[],
        )

        async def fake_run_scan(**_kwargs):
            return fake_scan_result

        with patch.object(routes_module, "run_scan", fake_run_scan):
            response = await self.client.post(
                "/api/apps/writing-review/scan",
                json={
                    "doc_text": "hello world",
                    "doc_name": "../../etc/passwd.md",
                },
            )
            self.assertEqual(response.status, 200)
            body = await response.json()
            job_id = body["job_id"]

            async with routes_module._JOBS_LOCK:
                job_record = dict(routes_module._JOBS[job_id])
            # Path traversal was stripped; ``..`` and ``/`` must not survive.
            self.assertEqual(job_record["doc_name"], "passwd.md")

            await asyncio.sleep(0.05)


class TestUploadRoute(_WritingReviewRoutesTestBase):
    """``POST /uploads`` accepts a raw file and writes bytes unchanged.

    The browse-file flow in ``NewReviewDialog`` used to route every
    browsed file through ``FileReader.readAsText`` and send the result
    in ``doc_text``. That works for ``.md`` / ``.txt`` but corrupts
    ``.docx`` -- a zipped bundle of XML -- because the JS string layer
    UTF-8-encodes the raw bytes and the on-disk copy is no longer a
    valid ZIP. ``parse_doc`` then hands it to python-docx and either
    fails or emits garbage sections.

    This endpoint is the frontend's escape hatch for binary browses.
    The bytes reach disk unmodified, the frontend receives a
    ``doc_path`` pointing at the stashed file, and the existing scan
    flow runs unchanged from there.
    """

    async def test_upload_writes_bytes_unchanged_and_returns_path(self) -> None:
        docx_first_bytes = b"PK\x03\x04"  # ZIP magic
        docx_bytes = docx_first_bytes + b"\x00\x01\x02" + b"\xff\xfe\xfd fake docx body"

        form_data = aiohttp.FormData()
        form_data.add_field(
            "file",
            docx_bytes,
            filename="my_design.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=form_data,
        )
        self.assertEqual(response.status, 200)
        response_body = await response.json()
        self.assertIn("doc_path", response_body)
        self.assertIn("doc_name", response_body)
        self.assertTrue(response_body["doc_path"].endswith(".docx"))

        # The bytes on disk MUST be identical to what we sent. This is
        # the whole point of the endpoint -- a byte-identical round-
        # trip so python-docx can open the ZIP later.
        stashed_bytes_on_disk = Path(response_body["doc_path"]).read_bytes()
        self.assertEqual(stashed_bytes_on_disk, docx_bytes)

    async def test_upload_rejects_missing_file_field(self) -> None:
        # An empty multipart body has no ``file`` field; the handler
        # must respond 400 with an ``error`` key rather than silently
        # writing a zero-byte artifact.
        empty_form = aiohttp.FormData()
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=empty_form,
        )
        self.assertEqual(response.status, 400)
        response_body = await response.json()
        self.assertIn("error", response_body)

    async def test_upload_rejects_empty_filename(self) -> None:
        # A multipart part with bytes but no ``filename=`` value cannot
        # be routed to the correct parser by extension. Falling back to
        # a default filename (``pasted.md``) would misroute docx bytes
        # to the markdown parser and produce garbage sections. Reject
        # with 400 so the caller re-uploads with a filename.
        docx_bytes = b"PK\x03\x04dummy"
        form_data = aiohttp.FormData()
        # ``filename=""`` produces a multipart part with an empty
        # ``filename=`` attribute, exercising the fallback path.
        form_data.add_field(
            "file",
            docx_bytes,
            filename="",
            content_type="application/octet-stream",
        )
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=form_data,
        )
        self.assertEqual(response.status, 400)
        response_body = await response.json()
        self.assertIn("filename", response_body.get("error", "").lower())

    async def test_upload_rejects_oversized_file(self) -> None:
        # Same 2 MiB cap the paste path enforces. Enforce at the
        # endpoint so a large upload cannot exhaust uploads dir bytes
        # before the ScanRequest even arrives.
        oversized_bytes = b"A" * (2 * 1024 * 1024 + 1)
        form_data = aiohttp.FormData()
        form_data.add_field(
            "file",
            oversized_bytes,
            filename="huge.docx",
            content_type="application/octet-stream",
        )
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=form_data,
        )
        self.assertEqual(response.status, 413)

    async def test_upload_sanitises_filename(self) -> None:
        # A browsed filename is user-supplied and never trusted --
        # the same sanitiser used by ``_stash_pasted_document`` for
        # the paste path applies here so path traversal cannot escape
        # the uploads directory.
        docx_bytes = b"PK\x03\x04dummy"
        form_data = aiohttp.FormData()
        form_data.add_field(
            "file",
            docx_bytes,
            filename="../../etc/passwd.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=form_data,
        )
        self.assertEqual(response.status, 200)
        response_body = await response.json()
        # Sanitiser reduces the browsed name to its basename and
        # replaces unsafe characters. The real security invariant is
        # that the stored file resolves inside the uploads directory
        # -- a filename like ``foo..bar.docx`` is fine as a substring
        # match on ``..``; a traversal only exists if a ``/`` or
        # os-specific separator escapes the intended parent.
        stored_path = Path(response_body["doc_path"])
        uploads_dir_expected = (
            Path(self._tempdir.name) / "apps" / "writing-review" / "uploads"
        ).resolve()
        self.assertEqual(stored_path.parent.resolve(), uploads_dir_expected)
        self.assertTrue(response_body["doc_path"].endswith(".docx"))

    async def test_upload_real_docx_round_trips_and_parses_body_correctly(self) -> None:
        """End-to-end validation with a real ``.docx`` on disk.

        Reads a hand-authored ``.docx`` (heading + one paragraph) from
        ``~/dev_specs/``, uploads it through the multipart endpoint,
        and then feeds the stashed copy to ``parse_doc``. Behavioural
        contract:

        * The bytes on disk match the source byte-for-byte (the whole
          reason ``/uploads`` exists -- ``readAsText`` on a ZIP would
          have re-encoded the bytes through UTF-8 and corrupted the
          archive).
        * ``parse_doc`` reads the stashed file cleanly (python-docx
          opens the ZIP).
        * The parsed section headings and body text match the source
          document exactly. This is what proves the upload path
          preserves not just the ZIP structure but every byte inside
          it -- a partial write or truncated body would still open in
          python-docx and might even parse without an error, but the
          extracted text would be wrong.

        Skipped in environments that do not have the fixture on disk.
        The fixture lives at ``~/dev_specs/test_reading_document.docx``
        on the maintainer's workstation and is intentionally NOT
        committed to the repo -- it is a local integration fixture,
        not a portable unit-test fixture.
        """
        real_docx_path = Path.home() / "dev_specs" / "test_reading_document.docx"
        if not real_docx_path.is_file():
            self.skipTest(f"integration fixture not present: {real_docx_path}")

        source_docx_bytes = real_docx_path.read_bytes()
        form_data = aiohttp.FormData()
        form_data.add_field(
            "file",
            source_docx_bytes,
            filename="test_reading_document.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
        response = await self.client.post(
            "/api/apps/writing-review/uploads",
            data=form_data,
        )
        self.assertEqual(response.status, 200)
        response_body = await response.json()

        stashed_bytes_on_disk = Path(response_body["doc_path"]).read_bytes()
        self.assertEqual(
            stashed_bytes_on_disk,
            source_docx_bytes,
            "uploaded .docx bytes must survive round-trip byte-for-byte",
        )

        # parse_doc is what the scan flow actually calls. Exercise the
        # real function against the real stashed file rather than a
        # mock -- this is what proves the upload endpoint delivered a
        # python-docx-openable archive.
        from kiro_crew.apps.builtins.writing_review import parse_doc

        parsed_sections = parse_doc(response_body["doc_path"])
        self.assertEqual(
            len(parsed_sections),
            1,
            "test fixture has exactly one heading + body pair",
        )
        self.assertEqual(parsed_sections[0].heading, "Test document")
        self.assertEqual(
            parsed_sections[0].body,
            "This is a test document to validate that the app has "
            "successfully read this document as expected.",
        )
