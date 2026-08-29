"""Unit tests for :func:`writing_review.parse_doc`.

Each ``TestCase`` maps to one row of the Behaviour Table in the spec:

* :class:`TestParseMarkdown`      -> Behaviour #1
* :class:`TestParsePlainText`     -> Behaviour #2
* :class:`TestParseDocx`          -> Behaviour #3
* :class:`TestSensitivePathGuard` -> Behaviour #4
* :class:`TestMissingFileGuard`   -> Behaviour #5
"""

from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from kiro_crew.apps.builtins.writing_review import parse_doc


class TestParseMarkdown(unittest.TestCase):
    """Behaviour #1 -- markdown documents split on ATX headings."""

    def setUp(self) -> None:
        self._tempdir = TemporaryDirectory()
        self.workspace_root = Path(self._tempdir.name)

    def tearDown(self) -> None:
        self._tempdir.cleanup()

    def test_splits_on_atx_headings(self) -> None:
        markdown_path = self.workspace_root / "notes.md"
        markdown_path.write_text("# Intro\nHello\n## Details\nWorld\n", encoding="utf-8")

        sections = parse_doc(markdown_path)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading, "Intro")
        self.assertEqual(sections[0].body, "Hello")
        self.assertEqual(sections[1].heading, "Details")
        self.assertEqual(sections[1].body, "World")

    def test_no_headings_returns_single_section(self) -> None:
        markdown_path = self.workspace_root / "flat.md"
        markdown_path.write_text("Just plain text\nwith multiple lines\n", encoding="utf-8")

        sections = parse_doc(markdown_path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading, "")
        self.assertIn("Just plain text", sections[0].body)


class TestParsePlainText(unittest.TestCase):
    """Behaviour #2 -- ``.txt`` files return a single heading-less section."""

    def setUp(self) -> None:
        self._tempdir = TemporaryDirectory()
        self.workspace_root = Path(self._tempdir.name)

    def tearDown(self) -> None:
        self._tempdir.cleanup()

    def test_returns_single_section(self) -> None:
        text_path = self.workspace_root / "notes.txt"
        text_path.write_text("Hello world\nLine two\n", encoding="utf-8")

        sections = parse_doc(text_path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading, "")
        self.assertIn("Hello world", sections[0].body)


class TestParseDocx(unittest.TestCase):
    """Behaviour #3 -- ``.docx`` files use Word heading styles for sections."""

    def setUp(self) -> None:
        self._tempdir = TemporaryDirectory()
        self.workspace_root = Path(self._tempdir.name)

    def tearDown(self) -> None:
        self._tempdir.cleanup()

    def test_extracts_heading_styled_paragraphs(self) -> None:
        from docx import Document as DocxDocument

        docx_path = self.workspace_root / "brief.docx"
        working_document = DocxDocument()
        working_document.add_heading("Summary", level=1)
        working_document.add_paragraph("Content here")
        working_document.save(str(docx_path))

        sections = parse_doc(docx_path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].heading, "Summary")
        self.assertIn("Content here", sections[0].body)

    def test_extracts_tables_as_markdown_pipe_tables(self) -> None:
        """A ``.docx`` with a table must surface the table content to the
        scanners. Rendered as a markdown pipe-table so the LLM can read
        the row/column structure -- otherwise scanners like structure
        Rule 5 (forward references to tables) false-positive on documents
        that DO have the referenced table right there. The table lands in
        document order under the current heading.
        """
        from docx import Document as DocxDocument

        docx_path = self.workspace_root / "with_table.docx"
        working_document = DocxDocument()
        working_document.add_heading("Fee Schedule", level=1)
        working_document.add_paragraph("The fees below apply from FY2025.")
        table_element = working_document.add_table(rows=3, cols=3)
        # Row 0 = header, rows 1 and 2 = body.
        header_cells = table_element.rows[0].cells
        header_cells[0].text = "Item"
        header_cells[1].text = "Quantity"
        header_cells[2].text = "MRC"
        body_row_one = table_element.rows[1].cells
        body_row_one[0].text = "Power"
        body_row_one[1].text = "52"
        body_row_one[2].text = "GBP 7020"
        body_row_two = table_element.rows[2].cells
        body_row_two[0].text = "ISP"
        body_row_two[1].text = "2"
        body_row_two[2].text = "GBP 840"
        working_document.add_paragraph("End of schedule.")
        working_document.save(str(docx_path))

        sections = parse_doc(docx_path)

        self.assertEqual(sections[0].heading, "Fee Schedule")
        # The three body pieces (paragraph, table, paragraph) must all
        # be present. Table content is markdown pipe-table shape.
        self.assertIn("The fees below apply from FY2025.", sections[0].body)
        self.assertIn("| Item | Quantity | MRC |", sections[0].body)
        self.assertIn("| --- | --- | --- |", sections[0].body)
        self.assertIn("| Power | 52 | GBP 7020 |", sections[0].body)
        self.assertIn("| ISP | 2 | GBP 840 |", sections[0].body)
        self.assertIn("End of schedule.", sections[0].body)

    def test_preserves_document_order_of_paragraphs_and_tables(self) -> None:
        """Doc order matters -- a table between two paragraphs must show
        up between them in the extracted body, not stapled onto the end.
        A structure scanner asked "does the fee reference in paragraph 2
        get followed by the fee schedule?" needs the table adjacent to
        its introducing sentence, not floating at the bottom.
        """
        from docx import Document as DocxDocument

        docx_path = self.workspace_root / "ordered.docx"
        working_document = DocxDocument()
        working_document.add_heading("Order", level=1)
        working_document.add_paragraph("BEFORE the table")
        interleaved_table = working_document.add_table(rows=1, cols=2)
        interleaved_table.rows[0].cells[0].text = "left"
        interleaved_table.rows[0].cells[1].text = "right"
        working_document.add_paragraph("AFTER the table")
        working_document.save(str(docx_path))

        sections = parse_doc(docx_path)

        body_text = sections[0].body
        before_index = body_text.find("BEFORE the table")
        table_index = body_text.find("| left | right |")
        after_index = body_text.find("AFTER the table")
        # All three markers present.
        self.assertGreaterEqual(before_index, 0)
        self.assertGreaterEqual(table_index, 0)
        self.assertGreaterEqual(after_index, 0)
        # And they appear in the source-document order.
        self.assertLess(before_index, table_index)
        self.assertLess(table_index, after_index)

    def test_image_becomes_a_placeholder_marker(self) -> None:
        """The scanners run text-only LLM prompts -- there is no vision
        path -- so an inline image cannot be reviewed for content. The
        extractor MUST still emit a placeholder so the LLM knows an
        image was PRESENT at that position. Otherwise a scanner
        reviewing paragraph 5 has no idea the paragraph referenced a
        figure and will emit a false 'missing figure' finding.
        """
        import io
        import struct
        import zlib

        from docx import Document as DocxDocument

        docx_path = self.workspace_root / "with_image.docx"
        working_document = DocxDocument()
        working_document.add_heading("Diagram section", level=1)
        working_document.add_paragraph("See the diagram below.")

        # Build the smallest valid PNG on the fly -- python-docx does not
        # ship an image factory, so we make a 1x1 pixel PNG in memory.
        # Constructed manually to avoid pulling Pillow into the test path.
        def _build_minimal_png_bytes() -> bytes:
            def _chunk(chunk_kind: bytes, chunk_data: bytes) -> bytes:
                data_len_prefix = struct.pack(">I", len(chunk_data))
                crc_bytes = struct.pack(">I", zlib.crc32(chunk_kind + chunk_data) & 0xFFFFFFFF)
                return data_len_prefix + chunk_kind + chunk_data + crc_bytes

            png_signature = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            raw_pixel_data = b"\x00\xff\xff\xff"  # one filter byte + one white pixel (RGB)
            idat_data = zlib.compress(raw_pixel_data)
            return (
                png_signature
                + _chunk(b"IHDR", ihdr_data)
                + _chunk(b"IDAT", idat_data)
                + _chunk(b"IEND", b"")
            )

        minimal_png_bytes = _build_minimal_png_bytes()
        working_document.add_picture(io.BytesIO(minimal_png_bytes))
        working_document.add_paragraph("Text after the diagram.")
        working_document.save(str(docx_path))

        sections = parse_doc(docx_path)

        body_text = sections[0].body
        # Placeholder MUST start with the strong VISUAL prefix and MUST
        # instruct the LLM to suppress its "missing diagram" finding on
        # the surrounding prose. The old ``[Image](Image exists...)``
        # wording was too weak -- scanners still flagged prose-heavy
        # sections for "missing diagram" even when an image was
        # embedded. The prefix + explicit suppression clause is the
        # signal that stops that false positive.
        self.assertIn("[VISUAL:", body_text)
        self.assertIn("Do NOT flag this section for missing", body_text)
        self.assertIn("embedded in the source document", body_text)
        self.assertIn("See the diagram below.", body_text)
        self.assertIn("Text after the diagram.", body_text)

    def test_image_placeholder_includes_alt_text_when_docpr_descr_present(self) -> None:
        """When the docx author added an accessibility ``descr`` on the
        image's ``<wp:docPr>``, the placeholder MUST embed that text as
        content evidence so the scanner reasons about what the image
        depicts rather than treating it as opaque. Alt-text is the
        author's own description of the visual and is the closest thing
        the text pipeline gets to seeing the image.
        """
        import io
        import struct
        import zlib

        from docx import Document as DocxDocument

        docx_path = self.workspace_root / "with_alt_text.docx"
        working_document = DocxDocument()
        working_document.add_heading("Network topology", level=1)

        def _build_minimal_png_bytes() -> bytes:
            def _chunk(chunk_kind: bytes, chunk_data: bytes) -> bytes:
                data_len_prefix = struct.pack(">I", len(chunk_data))
                crc_bytes = struct.pack(">I", zlib.crc32(chunk_kind + chunk_data) & 0xFFFFFFFF)
                return data_len_prefix + chunk_kind + chunk_data + crc_bytes

            png_signature = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            raw_pixel_data = b"\x00\xff\xff\xff"
            idat_data = zlib.compress(raw_pixel_data)
            return (
                png_signature
                + _chunk(b"IHDR", ihdr_data)
                + _chunk(b"IDAT", idat_data)
                + _chunk(b"IEND", b"")
            )

        minimal_png_bytes = _build_minimal_png_bytes()
        inline_shape = working_document.add_picture(io.BytesIO(minimal_png_bytes))
        # Set the descr on the <wp:docPr> element under the drawing.
        # python-docx exposes this via ``inline_shape._inline.docPr.set(...)``.
        drawing_element = inline_shape._inline
        drawing_element.docPr.set(
            "descr",
            "Network topology showing ISP to firewall to AGG to ToR flow",
        )
        working_document.save(str(docx_path))

        sections = parse_doc(docx_path)
        body_text = sections[0].body

        # Alt-text MUST appear in the placeholder so the scanner sees
        # what the image depicts and can reason about the surrounding
        # prose against the visual's actual content.
        self.assertIn(
            "Network topology showing ISP to firewall to AGG to ToR flow",
            body_text,
        )
        # Strong wording MUST still be present regardless of alt-text.
        self.assertIn("[VISUAL:", body_text)
        self.assertIn("Do NOT flag this section for missing", body_text)

    def test_real_amazon_order_form_docx_surfaces_fee_schedule_cells(self) -> None:
        """End-to-end integration on a real hand-authored ``.docx``.

        The Amazon Order Form has FOUR tables including the actual fee
        schedule (Item / QTY / MRC / NRC), which previously vanished
        entirely from the extractor's output. This test skips when the
        fixture isn't on disk so the test is portable, but pins the
        specific cell values a scanner needs to see to correctly reason
        about the doc's fee content.
        """
        real_docx_path = Path.home() / "dev_specs" / "amazon_order_form.docx"
        if not real_docx_path.is_file():
            self.skipTest(f"integration fixture not present: {real_docx_path}")

        sections = parse_doc(real_docx_path)
        combined_body = "\n".join(section.body for section in sections)
        # Fee schedule cell values -- a cell we KNOW is only in Table 1.
        self.assertIn("£135.00", combined_body)
        # Signature block cells -- Table 3.
        self.assertIn("Mark Pestridge", combined_body)
        # General terms cells -- Table 0.
        self.assertIn("BS88 HRC fuses", combined_body)


class TestSensitivePathGuard(unittest.TestCase):
    """Behaviour #4 -- reading a sensitive path raises ``PermissionError``."""

    def test_rejects_sensitive_path(self) -> None:
        # ``~/.aws/credentials`` is the canonical sensitive example and is
        # matched by ``kiro_crew.security.is_sensitive_path`` regardless of
        # whether the file actually exists on this host.
        sensitive_target = Path.home() / ".aws" / "credentials"
        with self.assertRaises(PermissionError):
            parse_doc(sensitive_target)


class TestMissingFileGuard(unittest.TestCase):
    """Behaviour #5 -- reading a missing path raises ``FileNotFoundError``."""

    def test_rejects_missing_file(self) -> None:
        with self.assertRaises(FileNotFoundError):
            parse_doc(Path("/nonexistent/writing-review/does-not-exist.md"))
